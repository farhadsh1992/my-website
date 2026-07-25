# -*- coding: utf-8 -*-
"""Generates an editable Word (.docx) version of the resume, matching
scripts/build_resume.py's content, so it can be redesigned by hand in
Microsoft Word.

    python3 scripts/build_resume_docx.py

Requires: pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "resume_source", "Farhad_Shadmand_Resume.docx")
OUT_PATH = os.path.abspath(OUT_PATH)
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)

PHOTO_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "public", "assets", "images", "profile.png"))

ACCENT = RGBColor(0xC0, 0x20, 0x2D)
INK = RGBColor(0x16, 0x18, 0x1D)
MUTED = RGBColor(0x5B, 0x5F, 0x6A)

doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = INK


def add_hyperlink(paragraph, url, text, color="1F5FBF", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), color)
        rPr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_cell_borders_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def add_bottom_border(paragraph, color="C0202D", size=18):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.color.rgb = ACCENT
    add_bottom_border(p)


def entry_header(title, dates):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.75), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    p.add_run("\t")
    r2 = p.add_run(dates)
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = MUTED
    return p


def meta_line(text, italic=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


def body_paragraph(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)
    return p


# ---------- Header ----------
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.LEFT
header_table.autofit = False
header_table.columns[0].width = Inches(1.9)
header_table.columns[1].width = Inches(4.65)

photo_cell, text_cell = header_table.rows[0].cells
photo_cell.width = Inches(1.9)
text_cell.width = Inches(4.65)
set_cell_borders_none(photo_cell)
set_cell_borders_none(text_cell)

if os.path.exists(PHOTO_PATH):
    p = photo_cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(PHOTO_PATH, width=Inches(1.75), height=Inches(1.68))

tp = text_cell.paragraphs[0]
tp.paragraph_format.space_after = Pt(0)
name_run = tp.add_run("Farhad Shadmand")
name_run.bold = True
name_run.font.size = Pt(22)
name_run.font.color.rgb = INK

title_p = text_cell.add_paragraph()
title_p.paragraph_format.space_before = Pt(1)
title_p.paragraph_format.space_after = Pt(4)
title_run = title_p.add_run("AI Researcher and Developer")
title_run.font.size = Pt(12)
title_run.font.color.rgb = MUTED

contact_items = [
    ("+351 912 292 634", None),
    ("farhadsh1992@gmail.com", "mailto:farhadsh1992@gmail.com"),
    ("www.farhadshad.com", "https://www.farhadshad.com"),
    ("linkedin.com/in/farhadsh1992", "https://www.linkedin.com/in/farhadsh1992/"),
    ("github.com/farhadsh1992", "https://www.github.com/farhadsh1992/"),
    ("orcid.org/0000-0003-4399-4845", "https://orcid.org/0000-0003-4399-4845"),
    ("ResearchGate", "https://www.researchgate.net/profile/Farhad-Shadmand-2"),
]
for label, url in contact_items:
    cp = text_cell.add_paragraph()
    cp.paragraph_format.space_after = Pt(0)
    if url:
        add_hyperlink(cp, url, label)
    else:
        r = cp.add_run(label)
        r.font.size = Pt(9.5)

hr = doc.add_paragraph()
hr.paragraph_format.space_before = Pt(6)
hr.paragraph_format.space_after = Pt(2)
add_bottom_border(hr, color="C0202D", size=18)

# ---------- About ----------
section_heading("About")
body_paragraph(
    "I am a researcher and developer with hands-on experience building deep learning models based on "
    "diffusion models, GANs, and VLM-based systems on multi-GPU platforms (CUDA). My work has led to "
    "multiple publications, including CVPR, WACV, and IEEE Access. I bring strong skills in PyTorch, "
    "computer vision, model optimization, and AI agents."
)

# ---------- Experience ----------
section_heading("Experience")

entry_header("Post-doctoral Researcher — VisTeam, Institute of Systems and Robotics – Coimbra (ISR)", "2026 – Present")
meta_line("Coimbra, Portugal")
bullet("Research on generative video models, including video watermarking and diffusion-based generation (Stable Diffusion).")
bullet("Stack: PyTorch, Hugging Face, Accelerate, Slurm, OpenCV, Weights & Biases (wandb).")

entry_header("Research Intern — Adobe", "Aug 2025 – Dec 2025")
meta_line("Paris, France")
bullet(
    "Developed robust watermarking models using geometric symmetry-aware convolutional networks combined with a "
    "VLM-based generative noise-simulation pipeline to accurately replicate print-and-scan degradation effects."
)
bullet("This approach significantly enhanced watermark resilience against real-world distortions introduced during printing, scanning, and image recapture.")
bullet("Stack: Deep Learning, AWS, PyTorch Lightning, data distribution.")
mgr_p = doc.add_paragraph()
mgr_p.paragraph_format.space_before = Pt(1)
mgr_p.paragraph_format.space_after = Pt(2)
mgr_run = mgr_p.add_run("Manager: Dr. Shruti Agarwal (")
mgr_run.italic = True
mgr_run.font.size = Pt(9.5)
mgr_run.font.color.rgb = MUTED
add_hyperlink(mgr_p, "mailto:shragarw@adobe.com", "shragarw@adobe.com")
mgr_run2 = mgr_p.add_run(")")
mgr_run2.italic = True
mgr_run2.font.size = Pt(9.5)
mgr_run2.font.color.rgb = MUTED

entry_header("Researcher — VisTeam, Institute of Systems and Robotics – Coimbra (ISR)", "Sep 2019 – Jul 2025")
meta_line("Coimbra, Portugal")
bullet("Developed deep learning models for steganography, watermarking, object detection, face detection, and face verification.")
bullet("Contributed to the projects VISUAL-ID, TruIM, and FACING, resulting in multiple publications and practical solutions for identity-document security.")
bullet("Stack: deep learning libraries, computer vision libraries (OpenCV, dlib, Pillow), C++, Django, FastAPI.")

entry_header("Researcher — Instituto Universitário de Lisboa (ISCTE)", "2018 – 2019")
meta_line("Lisbon, Portugal")
bullet("In 2018, I moved to Lisbon and began my research at Instituto Universitário de Lisboa (ISCTE) and the Faculty of Sciences of the University of Lisbon (FCUL), under the supervision of Prof. José Carlos Dias.")
bullet("Applied machine learning, deep learning, and natural language processing (NLP) techniques to analyze financial markets, processing data from online news websites, Twitter, and stock market platforms to develop predictive models for financial trends.")
bullet("Stack: NLP libraries, Python, MATLAB, Requests, BeautifulSoup, REST API.")

# ---------- Skills ----------
section_heading("Skills")
skill_groups = [
    ("Programming Languages", "Python, C++, Java (Basic), HTML, CSS"),
    ("AI/ML", "PyTorch, TensorFlow, ONNX, PyTorch Lightning, Hugging Face, DeepSpeed, Scikit-learn, "
              "OpenVINO, XGBoost, LightGBM, CatBoost, torchdata, Accelerate, Keras, JAX, FLAX, Optax, Core ML"),
    ("NLP", "spaCy, NLTK, AllenNLP, Transformers, Tokenizers, Optimum-Intel, Optimum-Neuron, "
            "Optimum-Habana, Optimum-ONNX, OpenAI SDK"),
    ("AI Agents", "LangChain, LangSmith, LangGraph, LangFlow, RAG, Gradio"),
    ("Systems", "Multi-GPU, CUDA, SLURM, Docker, Kubernetes (K8s), OpenHPC"),
    ("CI/CD", "Jenkins, GitHub Actions, GitLab CI"),
    ("Computer Vision", "OpenCV, dlib, Pillow"),
    ("Data", "SQL, MySQL, Spark, pandas, ChromaDB"),
]
for label, items in skill_groups:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = MUTED
    r2 = p.add_run(items)
    r2.font.size = Pt(10)

# ---------- Languages ----------
section_heading("Languages")
body_paragraph("Persian (Native), English (Fluent), Portuguese (Beginner)")

# ---------- Education ----------
section_heading("Education")

entry_header("PhD in Electrical Engineering and Intelligent Systems", "2021 – 2026")
org_p = doc.add_paragraph()
org_p.paragraph_format.space_after = Pt(1)
org_run = org_p.add_run("University of Coimbra — Coimbra, Portugal")
org_run.bold = True
org_run.font.size = Pt(9.5)
body_paragraph(
    "My PhD focused on developing printable steganography and watermarking models for identity-document "
    "security, combining encoder–decoder architectures, geometric-symmetric convolutions, and advanced "
    "noise-simulation pipelines. I designed and optimized several deep-learning frameworks for reliable "
    "message embedding under real print-and-capture distortions. The work resulted in multiple publications, "
    "including IEEE Access, CVPR, and WACV."
)
sup_p = doc.add_paragraph()
sup_p.paragraph_format.space_after = Pt(2)
sr = sup_p.add_run("Supervisors: Prof. Nuno Gonçalves (")
sr.italic = True
sr.font.size = Pt(9.5)
sr.font.color.rgb = MUTED
add_hyperlink(sup_p, "mailto:nunogon@deec.uc.pt", "nunogon@deec.uc.pt")
sr2 = sup_p.add_run(") and Prof. Luiz Schirmer (")
sr2.italic = True
sr2.font.size = Pt(9.5)
sr2.font.color.rgb = MUTED
add_hyperlink(sup_p, "mailto:Luiz.schirmer@ufsm.br", "Luiz.schirmer@ufsm.br")
sr3 = sup_p.add_run(")")
sr3.italic = True
sr3.font.size = Pt(9.5)
sr3.font.color.rgb = MUTED

entry_header("M.Sc. in Complex Systems, Physics", "2015 – 2017")
org_p2 = doc.add_paragraph()
org_p2.paragraph_format.space_after = Pt(1)
org_run2 = org_p2.add_run("Isfahan University of Technology — Isfahan, Iran")
org_run2.bold = True
org_run2.font.size = Pt(9.5)
body_paragraph(
    "I specialize in predicting and modeling stock market behavior, utilizing advanced data analysis and "
    "machine learning techniques to forecast trends and inform investment strategies."
)
sup2_p = doc.add_paragraph()
sup2_p.paragraph_format.space_after = Pt(2)
sr4 = sup2_p.add_run("Supervisor: Prof. Farhad Shahbazi (")
sr4.italic = True
sr4.font.size = Pt(9.5)
sr4.font.color.rgb = MUTED
add_hyperlink(sup2_p, "mailto:shahbazi@iut.ac.ir", "shahbazi@iut.ac.ir")
sr5 = sup2_p.add_run(")")
sr5.italic = True
sr5.font.size = Pt(9.5)
sr5.font.color.rgb = MUTED

entry_header("Bachelor's in Physics", "2011 – 2015")
org_p3 = doc.add_paragraph()
org_p3.paragraph_format.space_after = Pt(1)
org_run3 = org_p3.add_run("Isfahan University of Technology — Isfahan, Iran")
org_run3.bold = True
org_run3.font.size = Pt(9.5)
body_paragraph("Among the top 10% of graduates.")

# ---------- Patents ----------
section_heading("Patents")
entry_header(
    "Encoding, Decoding and Integrity Validation Systems for a Security Document with a Steganography-Encoded Image",
    "Granted Oct. 2025",
)
meta_line(
    "Inventors: Nuno Gonçalves and Farhad Shadmand · Applicants: Imprensa Nacional Casa da Moeda and University of Coimbra",
    italic=False,
)
body_paragraph(
    "Encoding, decoding and integrity validation systems for a security document with a steganography-encoded "
    "image and methods, computer programs and associated computer-readable data carrier."
)
patent_links_p = doc.add_paragraph()
patent_links_p.paragraph_format.space_after = Pt(2)
add_hyperlink(patent_links_p, "https://patents.google.com/patent/PT117136A/en", "PT117136A")
pl1 = patent_links_p.add_run(" (Portugal) · ")
pl1.font.size = Pt(9.5)
pl1.font.color.rgb = MUTED
add_hyperlink(patent_links_p, "https://patents.google.com/patent/EP4064095A1/en", "EP4064095A1")
pl2 = patent_links_p.add_run(" (European Patent Office)")
pl2.font.size = Pt(9.5)
pl2.font.color.rgb = MUTED

# ---------- Publications ----------
section_heading("Publications")
publications = [
    ("StampOne: Addressing Frequency Balance in Printer-proof Steganography",
     "CVPR (2024)",
     "https://openaccess.thecvf.com/content/CVPR2024W/WMF/html/Shadmand_StampOne_Addressing_Frequency_Balance_in_Printer-proof_Steganography_CVPRW_2024_paper.html"),
    ("CodeFace: A Deep Learning Printer-Proof Steganography for Face Portraits",
     "IEEE Access 9 (2021)",
     "https://ieeexplore.ieee.org/document/9634021/"),
    ("StylePuncher: Encoding a Hidden QR Code into Images",
     "ICPRAM (2025)",
     "https://www.scitepress.org/Papers/2025/131908/131908.pdf"),
    ("DocSafe: Towards Practical Print-Proof Image Steganography via Frequency Decomposition and Covariance Alignment",
     "IEEE Access (2026.3680290), 2026",
     "https://doi.org/10.1109/ACCESS.2026.3680290"),
    ("RiemStega: Covariance-based loss for print-proof transmission of data in images",
     "WACV (2025)",
     "https://openaccess.thecvf.com/content/WACV2025/papers/Cruz_RiemStega_Covariance-Based_Loss_for_Print-Proof_Transmission_of_Data_in_Images_WACV_2025_paper.pdf"),
    ("Young Labeled Faces in the Wild (YLFW): A Dataset for Children Faces Recognition",
     "IEEE 18th International Conference on Automatic Face and Gesture Recognition (FG) (2024)",
     "https://arxiv.org/abs/2301.05776"),
    ("MorDeephy: Face Morphing Detection via Fused Classification",
     "12th International Conference on Pattern Recognition Application and Methods (2022)",
     "https://arxiv.org/abs/2208.03110"),
    ("Towards Facial Biometrics for ID Document Validation in Mobile Devices",
     "Applied Sciences 11.13 (2021)",
     "https://doi.org/10.3390/app11136134"),
]
for title, venue, link in publications:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_hyperlink(p, link, title)
    vp = doc.add_paragraph()
    vp.paragraph_format.space_after = Pt(6)
    vr = vp.add_run(venue)
    vr.italic = True
    vr.font.size = Pt(9)
    vr.font.color.rgb = MUTED

doc.save(OUT_PATH)
print("wrote", OUT_PATH)
